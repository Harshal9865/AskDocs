import io
import uuid

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.models.document import Chunk, Document


IMAGE_TYPES = {"png", "jpg", "jpeg", "webp", "gif"}

MIME_MAP = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "gif": "image/gif",
}


def _pdf_has_text(data: bytes) -> bool:
    """Check if a PDF has extractable text (not just a scanned image)."""
    reader = PdfReader(io.BytesIO(data))
    for page in reader.pages[:5]:  # sample first 5 pages
        text = page.extract_text() or ""
        # consider it "has text" if there's meaningful content beyond whitespace
        stripped = text.strip()
        if len(stripped) > 50:
            return True
    return False


MAX_OCR_PAGES = 30


def _pdf_to_images(data: bytes, dpi: int = 150) -> list[bytes]:
    """Convert PDF pages to PNG images using pymupdf. Limited to MAX_OCR_PAGES."""
    try:
        import fitz  # pymupdf
    except ImportError as e:
        raise RuntimeError("pymupdf not installed — cannot OCR scanned PDFs") from e

    doc = fitz.open(stream=data, filetype="pdf")
    images: list[bytes] = []
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)
    for idx, page in enumerate(doc):
        if idx >= MAX_OCR_PAGES:
            break
        pix = page.get_pixmap(matrix=mat)
        images.append(pix.tobytes("png"))
    doc.close()
    return images


def extract_text(data: bytes, file_type: str) -> str:
    """Extract text from a document. For images and scanned PDFs, returns
    a sentinel value -- the caller must run Gemini OCR."""
    if file_type in IMAGE_TYPES:
        return ""  # signal: needs OCR
    if file_type == "pdf":
        reader = PdfReader(io.BytesIO(data))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text
        return ""  # signal: scanned PDF, needs OCR
    if file_type == "docx":
        doc = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return data.decode("utf-8", errors="replace")


_TOKEN_CHARS_PER_TOKEN = 4  # ~4 chars per token heuristic


def chunk_text(text: str) -> list[str]:
    """Token-approximate sliding-window chunking (~500 tokens, 50 overlap)."""
    settings = get_settings()
    size = settings.CHUNK_SIZE_TOKENS * _TOKEN_CHARS_PER_TOKEN
    overlap = settings.CHUNK_OVERLAP_TOKENS * _TOKEN_CHARS_PER_TOKEN

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    # build blocks of paragraphs up to size
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    def flush():
        nonlocal current, current_len
        if current:
            chunks.append("\n\n".join(current))
        current, current_len = [], 0

    for para in paragraphs:
        while len(para) > size:
            if current:
                flush()
            chunks.append(para[:size])
            para = para[size - overlap:]
        if current_len + len(para) > size and current:
            flush()
        current.append(para)
        current_len += len(para)
    flush()

    # dedupe tiny trailing chunk
    return [c for c in chunks if len(c) > 20]


async def _ocr_with_gemini(data: bytes, file_type: str) -> str:
    """Use Gemini vision to extract text from images or scanned PDFs."""
    import asyncio

    from app.services.llm.gemini_provider import get_llm

    llm = get_llm()

    if file_type in IMAGE_TYPES:
        mime = MIME_MAP.get(file_type, "image/png")
        try:
            return await llm.ocr_image(data, mime)
        except Exception:
            return ""

    if file_type == "pdf":
        try:
            images = _pdf_to_images(data)
        except Exception as e:
            raise RuntimeError(f"Failed to render PDF for OCR: {e}") from e
        # parallel OCR with limited concurrency (2 at a time to avoid rate limits)
        sem = asyncio.Semaphore(2)

        async def ocr_one(idx: int, img: bytes) -> str:
            async with sem:
                try:
                    text = await llm.ocr_image(img, "image/png")
                    return f"--- Page {idx + 1} ---\n{text}" if text.strip() else ""
                except Exception:
                    return ""

        results = await asyncio.gather(*(ocr_one(i, img) for i, img in enumerate(images)))
        combined = "\n\n".join(r for r in results if r.strip())
        if not combined.strip():
            raise ValueError("Gemini OCR returned no text — try a clearer scan or a text-based PDF")
        return combined

    return ""


async def ingest_document(db: AsyncSession, document_id: uuid.UUID):
    """Full pipeline: extract -> chunk -> embed -> store. Updates status."""
    document = await db.get(Document, document_id)
    if document is None:
        return
    from app.storage.db_storage import get_storage

    try:
        document.status = "processing"
        await db.commit()

        storage = get_storage()
        data = await storage.open(document.storage_key)
        text = extract_text(data, document.file_type)

        # fallback to Gemini OCR for images and scanned PDFs
        if not text.strip() and document.file_type in (IMAGE_TYPES | {"pdf"}):
            text = await _ocr_with_gemini(data, document.file_type)

        pieces = chunk_text(text)
        if not pieces:
            raise ValueError(
                "No extractable text found. "
                "The file may be empty or contain only images without text."
            )

        from app.services.llm.gemini_provider import get_llm

        llm = get_llm()
        embeddings = await llm.embed(pieces)

        for ordinal, (piece, embedding) in enumerate(zip(pieces, embeddings)):
            db.add(
                Chunk(
                    document_id=document.id,
                    workspace_id=document.workspace_id,
                    ordinal=ordinal,
                    content=piece,
                    token_count=len(piece) // _TOKEN_CHARS_PER_TOKEN,
                    metadata_json={},
                    embedding=embedding,
                )
            )
        document.status = "ready"
        document.error_msg = None
        await db.commit()
    except Exception as exc:  # noqa: BLE001
        await db.rollback()
        document = await db.get(Document, document_id)
        document.status = "failed"
        document.error_msg = str(exc)[:1000]
        await db.commit()
